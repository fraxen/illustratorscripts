# Convert old .doc metadata to XMP
import os
import tempfile
from docx import Document
from xml.etree import ElementTree

# First do this in c:\projection: attrib /s *.ai | cut -b14- > c:\home\hugo\work\allaifiles.txt
# Then we read that file into a list
"""
with open(r'c:\home\hugo\work\allaifiles.txt') as f:
    allAi = f.read()
allAi = allAi.split('\n')
"""

# Check where there are word files with the .ai files
"""
outFile = ''
for aifile in allAi:
    if os.path.exists(aifile[:-3] + '.doc'):
        outFile = outFile + aifile + '\n'

with open(r'c:\home\hugo\work\allaifiles2.txt','w') as f:
    f.write(outFile)
"""
with open(r'c:\home\hugo\work\allaifiles2.txt') as f:
    allMeta = f.read()
allMeta = allMeta.split('\n')
allMeta = filter(None, allMeta)

# Convert doc to docx so that we can interpret the files
import win32com.client
wrd = win32com.client.Dispatch("Word.Application")
wrd.visible = 0
for metaFile in allMeta:
    print '%s/%s - %s' % (allMeta.index(metaFile)+1, len(allMeta), metaFile)
    wb = wrd.Documents.Open(metaFile.replace('.ai', '.doc'))
    file_docx = tempfile.mktemp(suffix='.docx')
    wb.SaveAs2(file_docx, FileFormat=16)
    print 'converted: ' + metaFile
    wb.Close()

    # Read and interpret docx file
    document = Document(file_docx)
    fileMeta = {
        'title': '',
        'caption': '',
        'projection': '',
        'notes': '',
        'originator': 'Hugo Ahlenius, Nordpil',
        'credit': 'Hugo Ahlenius, Nordpil',
        'sources': []
    }
    para = document.paragraphs

    def getMeta(document, startSectionNum, outMeta):
        headings = ('title:', 'sources:', 'caption:', 'projection:', 'notes:',
                    'title', 'sources', 'caption', 'projection', 'notes')
        para = document.paragraphs
        for j in range(startSectionNum+1, len(para), 1):
            cleanedPara = para[j].text.replace(':', '').lower().strip()
            if cleanedPara in headings:
                break
            else:
                if type(outMeta) == list:
                    outMeta = outMeta + [para[j].text]
                else:
                    outMeta = (outMeta + '\n' if len(outMeta) else '') + para[j].text
        return outMeta

    for i in range(0, len(document.paragraphs), 1):
        cleanedPara = para[i].text.replace(':', '').lower().strip()
        if cleanedPara in fileMeta.keys():
            fileMeta[cleanedPara] = getMeta(document, i, fileMeta[cleanedPara])

    del document
    if fileMeta['title'] == '':
        fileMeta['title'] == os.path.basename(fileMeta)

    # Create XML file
    root = ElementTree.Element("graphicdoc")
    for el in ['title', 'caption', 'notes', 'projection', 'sources', 'credit', 'originator']:
        thisField = ElementTree.SubElement(root, el)
        if el == 'sources':
            for source in fileMeta['sources']:
                sourceField = ElementTree.SubElement(thisField, 'source')
                sourceField.text = source
        else:
            thisField.text = fileMeta[el]
    tree = ElementTree.ElementTree(root)
    tree.write(metaFile.replace('.ai', '.xml'))
wrd.Quit()
